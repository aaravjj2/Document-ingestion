#!/usr/bin/env python3
"""
Multi-Format Document Testing Script

Tests the Document Ingestion Service with:
1. PDF documents
2. Markdown files  
3. DOCX documents
4. PNG images
5. Handwritten content from the internet
"""

import asyncio
import sys
import time
import requests
from pathlib import Path
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import subprocess

print("=" * 80)
print("MULTI-FORMAT DOCUMENT INGESTION TEST")
print("=" * 80)
print()

# Create test directory
test_dir = Path("test_multi_format")
test_dir.mkdir(exist_ok=True)

# ============================================================================
# PART 1: Create Test Documents in Various Formats
# ============================================================================

print("📝 Creating test documents in multiple formats...")
print()

# 1. Create PNG Invoice
def create_png_invoice():
    """Create a PNG invoice."""
    print("  → Creating PNG invoice...")
    img = Image.new('RGB', (800, 1000), color='white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
        font_med = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
    except:
        font_large = font_med = font_small = ImageFont.load_default()
    
    # Header
    draw.rectangle([0, 0, 800, 80], fill='#2C3E50')
    draw.text((50, 25), "INVOICE", fill='white', font=font_large)
    
    # Invoice details
    draw.text((50, 120), "Invoice Number: INV-2024-12345", fill='black', font=font_med)
    draw.text((50, 160), "Date: December 30, 2024", fill='black', font=font_small)
    draw.text((50, 190), "Due Date: January 30, 2025", fill='black', font=font_small)
    
    # Bill to
    draw.text((50, 250), "Bill To:", fill='black', font=font_med)
    draw.text((50, 285), "Acme Corporation", fill='black', font=font_small)
    draw.text((50, 310), "123 Business Street", fill='black', font=font_small)
    draw.text((50, 335), "San Francisco, CA 94105", fill='black', font=font_small)
    
    # Items
    draw.line([50, 400, 750, 400], fill='black', width=2)
    draw.text((50, 420), "Description", fill='black', font=font_med)
    draw.text((500, 420), "Quantity", fill='black', font=font_med)
    draw.text((650, 420), "Amount", fill='black', font=font_med)
    draw.line([50, 450, 750, 450], fill='black', width=1)
    
    draw.text((50, 470), "Professional Services", fill='black', font=font_small)
    draw.text((520, 470), "10", fill='black', font=font_small)
    draw.text((660, 470), "$1,500.00", fill='black', font=font_small)
    
    draw.text((50, 510), "Software License", fill='black', font=font_small)
    draw.text((520, 510), "5", fill='black', font=font_small)
    draw.text((660, 510), "$2,500.00", fill='black', font=font_small)
    
    draw.line([50, 570, 750, 570], fill='black', width=1)
    
    # Totals
    draw.text((500, 600), "Subtotal:", fill='black', font=font_med)
    draw.text((660, 600), "$4,000.00", fill='black', font=font_med)
    
    draw.text((500, 640), "Tax (8.5%):", fill='black', font=font_small)
    draw.text((660, 640), "$340.00", fill='black', font=font_small)
    
    draw.line([500, 670, 750, 670], fill='black', width=2)
    draw.text((500, 690), "Total:", fill='black', font=font_large)
    draw.text((650, 690), "$4,340.00", fill='black', font=font_large)
    
    # Footer
    draw.text((50, 850), "Payment Terms: Net 30", fill='#7F8C8D', font=font_small)
    draw.text((50, 880), "Thank you for your business!", fill='#7F8C8D', font=font_small)
    
    path = test_dir / "invoice.png"
    img.save(path)
    print(f"    ✓ Created: {path}")
    return path

# 2. Create PDF Invoice
def create_pdf_invoice():
    """Create a PDF invoice using reportlab."""
    print("  → Creating PDF invoice...")
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        path = test_dir / "invoice.pdf"
        c = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        
        # Header
        c.setFont("Helvetica-Bold", 24)
        c.drawString(50, height - 50, "INVOICE")
        
        # Invoice details
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 100, "Invoice Number: INV-2024-67890")
        c.drawString(50, height - 120, "Date: December 30, 2024")
        c.drawString(50, height - 140, "Due Date: January 30, 2025")
        
        # Bill to
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 180, "Bill To:")
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 200, "TechStart Inc.")
        c.drawString(50, height - 220, "456 Innovation Drive")
        c.drawString(50, height - 240, "Austin, TX 78701")
        
        # Items table
        y = height - 300
        c.line(50, y, 550, y)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y - 20, "Description")
        c.drawString(350, y - 20, "Quantity")
        c.drawString(470, y - 20, "Amount")
        
        y -= 40
        c.line(50, y, 550, y)
        
        c.setFont("Helvetica", 11)
        y -= 25
        c.drawString(50, y, "Consulting Services")
        c.drawString(370, y, "20")
        c.drawString(470, y, "$3,000.00")
        
        y -= 25
        c.drawString(50, y, "Cloud Infrastructure")
        c.drawString(370, y, "1")
        c.drawString(470, y, "$1,200.00")
        
        y -= 40
        c.line(50, y, 550, y)
        
        # Totals
        y -= 30
        c.setFont("Helvetica-Bold", 12)
        c.drawString(350, y, "Subtotal:")
        c.drawString(470, y, "$4,200.00")
        
        y -= 25
        c.setFont("Helvetica", 11)
        c.drawString(350, y, "Tax (6%):")
        c.drawString(470, y, "$252.00")
        
        y -= 30
        c.line(350, y, 550, y)
        
        y -= 30
        c.setFont("Helvetica-Bold", 14)
        c.drawString(350, y, "Total:")
        c.drawString(460, y, "$4,452.00")
        
        # Footer
        c.setFont("Helvetica", 10)
        c.drawString(50, 50, "Payment Terms: Net 30 | Thank you for your business!")
        
        c.save()
        print(f"    ✓ Created: {path}")
        return path
    except ImportError:
        print("    ⚠️  reportlab not installed, skipping PDF creation")
        return None

# 3. Create Markdown Receipt
def create_markdown_receipt():
    """Create a markdown receipt."""
    print("  → Creating Markdown receipt...")
    content = """# RECEIPT

**Store:** QuickMart Express  
**Date:** December 30, 2024  
**Time:** 14:35:22  
**Transaction ID:** TXN-2024-98765

---

## Items Purchased

| Item | Quantity | Price |
|------|----------|-------|
| Organic Coffee Beans | 2 | $18.99 |
| Whole Milk (1 Gallon) | 1 | $4.49 |
| Fresh Bread | 3 | $8.97 |
| Orange Juice | 1 | $6.99 |

---

**Subtotal:** $39.44  
**Tax (7%):** $2.76  
**Total:** $42.20

**Payment Method:** Visa ending in 4532  
**Authorization:** APPROVED

---

*Thank you for shopping with us!*  
*Visit us again soon!*

Customer Service: 1-800-QUICKMART
"""
    path = test_dir / "receipt.md"
    path.write_text(content)
    print(f"    ✓ Created: {path}")
    return path

# 4. Create DOCX Document
def create_docx_medical():
    """Create a DOCX medical document."""
    print("  → Creating DOCX medical record...")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        
        # Title
        title = doc.add_heading('MEDICAL RECORD', 0)
        title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Patient Info
        doc.add_heading('Patient Information', level=1)
        doc.add_paragraph('Name: John Smith')
        doc.add_paragraph('Date of Birth: January 15, 1985')
        doc.add_paragraph('Patient ID: PT-2024-12345')
        doc.add_paragraph('Date: December 30, 2024')
        
        # Visit Info
        doc.add_heading('Visit Information', level=1)
        doc.add_paragraph('Chief Complaint: Routine checkup')
        doc.add_paragraph('Physician: Dr. Sarah Johnson, MD')
        doc.add_paragraph('Blood Pressure: 120/80 mmHg')
        doc.add_paragraph('Temperature: 98.6°F')
        doc.add_paragraph('Heart Rate: 72 bpm')
        
        # Diagnosis
        doc.add_heading('Diagnosis', level=1)
        doc.add_paragraph('Patient is in good health. No significant issues identified.')
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph('1. Continue current exercise routine')
        doc.add_paragraph('2. Maintain balanced diet')
        doc.add_paragraph('3. Schedule follow-up in 6 months')
        
        # Medications
        doc.add_heading('Medications', level=1)
        doc.add_paragraph('None prescribed at this time')
        
        # Footer
        doc.add_paragraph('')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Physician Signature: Dr. Sarah Johnson')
        doc.add_paragraph('Date: December 30, 2024')
        
        path = test_dir / "medical_record.docx"
        doc.save(str(path))
        print(f"    ✓ Created: {path}")
        return path
    except ImportError:
        print("    ⚠️  python-docx not installed, skipping DOCX creation")
        return None

# 5. Download handwritten content samples
def download_handwritten_samples():
    """Download handwritten text samples from the internet."""
    print("  → Downloading handwritten samples from internet...")
    
    samples = []
    
    # Sample URLs of handwritten text (these are publicly available test images)
    urls = [
        # IAM Handwriting Database samples
        ("https://fki.tic.heia-fr.ch/static/img/a01-000u-s00-00.png", "handwritten_form.png"),
        # MNIST-like handwritten digits
        ("https://upload.wikimedia.org/wikipedia/commons/2/27/MnistExamples.png", "handwritten_digits.png"),
        # Handwritten note sample
        ("https://raw.githubusercontent.com/sueiras/handwritting_recognition_CNN/master/data/words/r06-022-03-05.png", "handwritten_word.png"),
    ]
    
    for url, filename in urls:
        try:
            print(f"    • Downloading {filename}...")
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                path = test_dir / filename
                path.write_bytes(response.content)
                samples.append(path)
                print(f"      ✓ Downloaded: {path}")
            else:
                print(f"      ⚠️  Failed to download (status {response.status_code})")
        except Exception as e:
            print(f"      ⚠️  Error: {str(e)[:50]}")
    
    # Create our own handwritten-style sample
    print("    • Creating synthetic handwritten sample...")
    try:
        img = Image.new('RGB', (600, 400), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a script-like font
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSerif-Italic.ttf", 28)
        except:
            font = ImageFont.load_default()
        
        # Write in cursive-like style
        text = "Dear Customer,\n\nThank you for your order!\nYour package will arrive soon.\n\nBest regards,\nCustomer Service"
        
        y = 50
        for line in text.split('\n'):
            draw.text((50, y), line, fill='#1a1a1a', font=font)
            y += 40
        
        path = test_dir / "handwritten_note.png"
        img.save(path)
        samples.append(path)
        print(f"      ✓ Created: {path}")
    except Exception as e:
        print(f"      ⚠️  Error creating sample: {e}")
    
    return samples

# Create all test documents
png_invoice = create_png_invoice()
pdf_invoice = create_pdf_invoice()
md_receipt = create_markdown_receipt()
docx_medical = create_docx_medical()
handwritten = download_handwritten_samples()

print()
print(f"✓ Created {sum(1 for x in [png_invoice, pdf_invoice, md_receipt, docx_medical] if x)} test documents")
print(f"✓ Downloaded/Created {len(handwritten)} handwritten samples")
print()

# ============================================================================
# PART 2: Wait for Services to be Ready
# ============================================================================

print("=" * 80)
print("WAITING FOR SERVICES")
print("=" * 80)
print()

def wait_for_service(url, name, max_retries=60):
    """Wait for a service to be ready."""
    print(f"Waiting for {name}...")
    for i in range(max_retries):
        try:
            response = requests.get(url, timeout=2)
            if response.status_code < 500:
                print(f"  ✓ {name} is ready!")
                return True
        except:
            pass
        
        if i % 10 == 0 and i > 0:
            print(f"  ... still waiting ({i}/{max_retries})")
        time.sleep(1)
    
    print(f"  ✗ {name} failed to start")
    return False

# Check if API is running
api_ready = wait_for_service("http://localhost:8000/api/v1/dashboard/health", "API")

if not api_ready:
    print()
    print("⚠️  API is not running. Please start it first:")
    print("   docker-compose up -d")
    print()
    sys.exit(1)

print()

# ============================================================================
# PART 3: Test Document Upload and Processing
# ============================================================================

print("=" * 80)
print("TESTING DOCUMENT UPLOAD & PROCESSING")
print("=" * 80)
print()

test_files = [
    (png_invoice, "PNG Invoice"),
    (pdf_invoice, "PDF Invoice"),
    (md_receipt, "Markdown Receipt"),
    (docx_medical, "DOCX Medical Record"),
]

# Add handwritten samples
for hw_path in handwritten:
    test_files.append((hw_path, f"Handwritten: {hw_path.name}"))

results = []

for file_path, description in test_files:
    if file_path is None or not file_path.exists():
        print(f"⊘ Skipping {description} (not created)")
        continue
    
    print(f"📤 Testing: {description}")
    print(f"   File: {file_path}")
    
    try:
        # Upload document
        with open(file_path, 'rb') as f:
            files = {'file': (file_path.name, f, 'application/octet-stream')}
            response = requests.post('http://localhost:8000/api/v1/documents/upload', files=files, timeout=30)
        
        if response.status_code in (200, 202):  # Accept both 200 and 202 Accepted
            data = response.json()
            doc_id = data.get('document_id')
            job_id = data.get('job_id')
            
            print(f"   ✓ Uploaded successfully")
            print(f"   • Document ID: {doc_id}")
            print(f"   • Job ID: {job_id}")
            
            # Wait a bit for processing
            time.sleep(3)
            
            # Check status
            status_response = requests.get(f'http://localhost:8000/api/v1/documents/{doc_id}/status', timeout=10)
            if status_response.status_code == 200:
                status_data = status_response.json()
                print(f"   • Status: {status_data.get('status')}")
                print(f"   • Type: {status_data.get('document_type', 'unknown')}")
                
                # Get full details
                detail_response = requests.get(f'http://localhost:8000/api/v1/documents/{doc_id}', timeout=10)
                if detail_response.status_code == 200:
                    details = detail_response.json()
                    raw_text = details.get('raw_text', '')
                    confidence = details.get('ocr_confidence', 0)
                    
                    print(f"   • OCR Confidence: {confidence:.1%}" if confidence else "   • OCR Confidence: N/A")
                    
                    if raw_text:
                        preview = raw_text[:150].replace('\n', ' ')
                        print(f"   • Text Preview: {preview}...")
                    
                    results.append({
                        'description': description,
                        'doc_id': doc_id,
                        'status': status_data.get('status'),
                        'type': status_data.get('document_type'),
                        'confidence': confidence,
                        'text_length': len(raw_text) if raw_text else 0
                    })
            
            print(f"   ✓ Test completed")
        else:
            print(f"   ✗ Upload failed: {response.status_code}")
            print(f"      {response.text[:200]}")
    
    except Exception as e:
        print(f"   ✗ Error: {str(e)[:100]}")
    
    print()

# ============================================================================
# PART 4: Summary Report
# ============================================================================

print("=" * 80)
print("TEST SUMMARY")
print("=" * 80)
print()

print(f"Total Documents Tested: {len(results)}")
print()

if results:
    print("┌" + "─" * 78 + "┐")
    print("│ {:30} │ {:10} │ {:12} │ {:8} │".format("Description", "Status", "Type", "Text Len"))
    print("├" + "─" * 78 + "┤")
    
    for result in results:
        desc = result['description'][:30]
        status = result['status'][:10]
        doc_type = str(result.get('type', 'N/A'))[:12]
        text_len = result.get('text_length', 0)
        
        print("│ {:30} │ {:10} │ {:12} │ {:8} │".format(desc, status, doc_type, text_len))
    
    print("└" + "─" * 78 + "┘")
    print()
    
    # Statistics
    completed = sum(1 for r in results if r['status'] == 'completed')
    processing = sum(1 for r in results if r['status'] == 'processing')
    failed = sum(1 for r in results if r['status'] == 'failed')
    
    print("Statistics:")
    print(f"  ✓ Completed: {completed}/{len(results)}")
    print(f"  ⏳ Processing: {processing}/{len(results)}")
    print(f"  ✗ Failed: {failed}/{len(results)}")
    print()
    
    # Format breakdown
    format_counts = {}
    for result in results:
        fmt = result['description'].split(':')[0] if ':' in result['description'] else result['description']
        format_counts[fmt] = format_counts.get(fmt, 0) + 1
    
    print("Format Coverage:")
    for fmt, count in format_counts.items():
        print(f"  • {fmt}: {count} document(s)")

print()
print("=" * 80)
print("MULTI-FORMAT TEST COMPLETE!")
print("=" * 80)
print()
print("Test files saved in:", test_dir.absolute())
print("View results in dashboard: http://localhost:8501")
print()
